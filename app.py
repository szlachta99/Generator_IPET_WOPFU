import customtkinter as ctk
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import messagebox, filedialog
import os
from datetime import datetime

class EduApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        self.title("Generator IPET i WOPFU - System Wspomagania Nauczyciela")
        self.geometry("1000x750")
        ctk.set_appearance_mode("light")
        
        # Domyślna ścieżka zapisu (katalog aplikacji)
        self.save_directory = os.path.dirname(os.path.abspath(__file__))
        
        # Nawigacja boczna
        self.navigation_frame = ctk.CTkFrame(self, corner_radius=0, width=220)
        self.navigation_frame.grid(row=0, column=0, sticky="nsew")
        self.navigation_frame.grid_propagate(False)
        
        self.label_title = ctk.CTkLabel(self.navigation_frame, text="Panel Sterowania", 
                                        font=ctk.CTkFont(size=20, weight="bold"))
        self.label_title.pack(padx=20, pady=20)
        
        self.btn_wopfu = ctk.CTkButton(self.navigation_frame, text="Formularz WOPFU", 
                                       command=self.show_wopfu)
        self.btn_wopfu.pack(padx=20, pady=10)
        
        self.btn_ipet = ctk.CTkButton(self.navigation_frame, text="Formularz IPET", 
                                      command=self.show_ipet)
        self.btn_ipet.pack(padx=20, pady=10)
        
        # Przycisk wyboru ścieżki zapisu
        self.btn_choose_path = ctk.CTkButton(self.navigation_frame, 
                                            text="Wybierz folder zapisu",
                                            command=self.choose_save_directory,
                                            fg_color="gray")
        self.btn_choose_path.pack(padx=20, pady=10)
        
        # Etykieta z aktualną ścieżką
        self.label_current_path = ctk.CTkLabel(self.navigation_frame, 
                                              text=f"Zapisuję do:\n{self.get_short_path()}",
                                              font=ctk.CTkFont(size=10),
                                              wraplength=180)
        self.label_current_path.pack(padx=20, pady=(0, 20))
        
        # Kontener na formularze
        self.main_frame = ctk.CTkScrollableFrame(self, corner_radius=0, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        self.show_wopfu()
    
    def get_short_path(self):
        """Skraca długą ścieżkę dla lepszego wyświetlania"""
        path = self.save_directory
        if len(path) > 30:
            return "..." + path[-27:]
        return path
    
    def choose_save_directory(self):
        """Pozwala użytkownikowi wybrać folder do zapisu plików"""
        new_dir = filedialog.askdirectory(
            title="Wybierz folder do zapisu dokumentów",
            initialdir=self.save_directory
        )
        
        if new_dir:
            self.save_directory = new_dir
            self.label_current_path.configure(text=f"Zapisuję do:\n{self.get_short_path()}")
            messagebox.showinfo("Sukces", f"Folder zapisu zmieniony na:\n{new_dir}")
    
    def clear_main_frame(self):
        for widget in self.main_frame.winfo_children():
            widget.destroy()
    
    def show_wopfu(self):
        self.clear_main_frame()
        
        header = ctk.CTkLabel(self.main_frame, 
                            text="Wielospecjalistyczna Ocena Poziomu Funkcjonowania Ucznia (WOPFU)",
                            font=("Arial", 16, "bold"))
        header.pack(pady=10)
        
        self.wopfu_student_name = self.create_input("Imię i nazwisko ucznia:")
        self.wopfu_birth_date = self.create_input("Data urodzenia:")
        self.wopfu_school = self.create_input("Nazwa szkoły/przedszkola:")
        self.wopfu_class = self.create_input("Klasa/grupa:")
        
        self.wopfu_strengths = self.create_input("Mocne strony i zasoby ucznia:", area=True, height=80)
        self.wopfu_difficulties = self.create_input("Trudności i bariery:", area=True, height=80)
        self.wopfu_recommendations = self.create_input("Zalecenia dotyczące wsparcia:", area=True, height=80)
        
        btn_save = ctk.CTkButton(self.main_frame, text="Generuj Dokument WOPFU", 
                                fg_color="green", command=self.generate_wopfu_docx)
        btn_save.pack(pady=20)
    
    def create_input(self, label_text, area=False, height=100):
        label = ctk.CTkLabel(self.main_frame, text=label_text, anchor="w")
        label.pack(anchor="w", padx=20, pady=(10, 5))
        
        if area:
            entry = ctk.CTkTextbox(self.main_frame, height=height, width=700)
        else:
            entry = ctk.CTkEntry(self.main_frame, width=700)
        
        entry.pack(padx=20, pady=(0, 5))
        return entry
    
    def show_ipet(self):
        self.clear_main_frame()
        
        header = ctk.CTkLabel(self.main_frame, 
                            text="Indywidualny Program Edukacyjno-Terapeutyczny (IPET)",
                            font=("Arial", 16, "bold"))
        header.pack(pady=10)
        
        # Dane podstawowe
        section_label = ctk.CTkLabel(self.main_frame, text="DANE PODSTAWOWE", 
                                    font=("Arial", 14, "bold"), text_color="blue")
        section_label.pack(anchor="w", padx=20, pady=(15, 5))
        
        self.ipet_student_name = self.create_input("Imię i nazwisko ucznia:")
        self.ipet_birth_date = self.create_input("Data urodzenia:")
        self.ipet_school = self.create_input("Nazwa szkoły/przedszkola:")
        self.ipet_class = self.create_input("Klasa/grupa:")
        self.ipet_certificate_number = self.create_input("Numer orzeczenia o potrzebie kształcenia specjalnego:")
        self.ipet_certificate_date = self.create_input("Data wydania orzeczenia:")
        self.ipet_valid_until = self.create_input("Ważność orzeczenia do:")
        
        # WOPFU - podstawa do IPET
        section_label = ctk.CTkLabel(self.main_frame, text="OCENA FUNKCJONOWANIA (WOPFU)", 
                                    font=("Arial", 14, "bold"), text_color="blue")
        section_label.pack(anchor="w", padx=20, pady=(15, 5))
        
        self.ipet_strengths = self.create_input("Mocne strony ucznia:", area=True, height=70)
        self.ipet_difficulties = self.create_input("Trudności i ograniczenia:", area=True, height=70)
        self.ipet_needs = self.create_input("Potrzeby rozwojowe i edukacyjne:", area=True, height=70)
        
        # Cele edukacyjno-terapeutyczne
        section_label = ctk.CTkLabel(self.main_frame, text="CELE EDUKACYJNO-TERAPEUTYCZNE", 
                                    font=("Arial", 14, "bold"), text_color="blue")
        section_label.pack(anchor="w", padx=20, pady=(15, 5))
        
        self.ipet_educational_goals = self.create_input("Cele edukacyjne (np. opanowanie umiejętności czytania, pisania):", 
                                                        area=True, height=80)
        self.ipet_therapeutic_goals = self.create_input("Cele terapeutyczne (np. rozwój kompetencji społecznych, komunikacyjnych):", 
                                                        area=True, height=80)
        self.ipet_social_goals = self.create_input("Cele wychowawcze i społeczne:", area=True, height=60)
        
        # Dostosowania
        section_label = ctk.CTkLabel(self.main_frame, text="DOSTOSOWANIA I WSPARCIE", 
                                    font=("Arial", 14, "bold"), text_color="blue")
        section_label.pack(anchor="w", padx=20, pady=(15, 5))
        
        self.ipet_adjustments = self.create_input("Zakres dostosowań wymagań edukacyjnych (np. wydłużony czas, zmniejszona objętość materiału):", 
                                                  area=True, height=80)
        self.ipet_support_forms = self.create_input("Formy wsparcia (zajęcia rewalidacyjne, logopeda, psycholog, terapeuta SI itp.):", 
                                                    area=True, height=80)
        self.ipet_teaching_methods = self.create_input("Metody i sposoby pracy (AAC, technologie wspomagające, metody wizualne itp.):", 
                                                       area=True, height=80)
        
        # Zajęcia dodatkowe
        section_label = ctk.CTkLabel(self.main_frame, text="ZAJĘCIA SPECJALISTYCZNE", 
                                    font=("Arial", 14, "bold"), text_color="blue")
        section_label.pack(anchor="w", padx=20, pady=(15, 5))
        
        self.ipet_therapy_sessions = self.create_input("Zajęcia terapeutyczne (osoba, godziny tygodniowo, tematyka):", 
                                                       area=True, height=100)
        
        # Ewaluacja
        section_label = ctk.CTkLabel(self.main_frame, text="EWALUACJA", 
                                    font=("Arial", 14, "bold"), text_color="blue")
        section_label.pack(anchor="w", padx=20, pady=(15, 5))
        
        self.ipet_evaluation = self.create_input("Sposoby i terminy ewaluacji (min. 2 razy w roku):", 
                                                 area=True, height=60)
        self.ipet_team_members = self.create_input("Członkowie zespołu opracowującego IPET (wychowawca, pedagog specjalny, psycholog itp.):", 
                                                   area=True, height=60)
        
        btn_save = ctk.CTkButton(self.main_frame, text="Generuj Dokument IPET", 
                                fg_color="green", command=self.generate_ipet_docx)
        btn_save.pack(pady=30)
    
    def generate_wopfu_docx(self):
        # Pobieranie danych z UI
        name = self.wopfu_student_name.get().strip()
        birth_date = self.wopfu_birth_date.get().strip()
        school = self.wopfu_school.get().strip()
        class_name = self.wopfu_class.get().strip()
        
        str_val = self.wopfu_strengths.get("1.0", "end").strip()
        diff_val = self.wopfu_difficulties.get("1.0", "end").strip()
        rec_val = self.wopfu_recommendations.get("1.0", "end").strip()
        
        # Walidacja
        if not name:
            messagebox.showwarning("Błąd", "Proszę podać imię i nazwisko ucznia!")
            return
        
        # Tworzenie dokumentu
        doc = Document()
        
        # Nagłówek
        title = doc.add_heading('WIELOSPECJALISTYCZNA OCENA POZIOMU FUNKCJONOWANIA UCZNIA', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph(f"Data opracowania: {datetime.now().strftime('%d.%m.%Y')}")
        doc.add_paragraph()
        
        # Dane ucznia
        doc.add_heading('DANE UCZNIA', level=1)
        doc.add_paragraph(f"Imię i nazwisko: {name}")
        if birth_date:
            doc.add_paragraph(f"Data urodzenia: {birth_date}")
        if school:
            doc.add_paragraph(f"Szkoła/Przedszkole: {school}")
        if class_name:
            doc.add_paragraph(f"Klasa/Grupa: {class_name}")
        
        doc.add_paragraph()
        
        # Ocena funkcjonowania
        doc.add_heading('MOCNE STRONY I ZASOBY UCZNIA', level=1)
        doc.add_paragraph(str_val if str_val else "Nie podano.")
        
        doc.add_heading('TRUDNOŚCI I BARIERY', level=1)
        doc.add_paragraph(diff_val if diff_val else "Nie podano.")
        
        doc.add_heading('ZALECENIA DOTYCZĄCE WSPARCIA', level=1)
        doc.add_paragraph(rec_val if rec_val else "Nie podano.")
        
        # Zapisywanie
        safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_')).replace(' ', '_')
        filename = f"WOPFU_{safe_name}.docx"
        filepath = os.path.join(self.save_directory, filename)
        
        try:
            doc.save(filepath)
            messagebox.showinfo("Sukces", f"Dokument został zapisany:\n{filepath}")
        except PermissionError:
            messagebox.showerror("Błąd", 
                               "Brak uprawnień do zapisu w tej lokalizacji lub plik jest otwarty!\n"
                               "Spróbuj wybrać inny folder lub zamknij plik jeśli jest otwarty.")
        except Exception as e:
            messagebox.showerror("Błąd", f"Wystąpił błąd podczas zapisu:\n{str(e)}")
    
    def generate_ipet_docx(self):
        # Pobieranie danych podstawowych
        name = self.ipet_student_name.get().strip()
        birth_date = self.ipet_birth_date.get().strip()
        school = self.ipet_school.get().strip()
        class_name = self.ipet_class.get().strip()
        cert_number = self.ipet_certificate_number.get().strip()
        cert_date = self.ipet_certificate_date.get().strip()
        valid_until = self.ipet_valid_until.get().strip()
        
        # Walidacja
        if not name:
            messagebox.showwarning("Błąd", "Proszę podać imię i nazwisko ucznia!")
            return
        if not cert_number:
            messagebox.showwarning("Błąd", "Proszę podać numer orzeczenia!")
            return
        
        # Tworzenie dokumentu
        doc = Document()
        
        # Nagłówek
        title = doc.add_heading('INDYWIDUALNY PROGRAM EDUKACYJNO-TERAPEUTYCZNY', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph(f"Data opracowania: {datetime.now().strftime('%d.%m.%Y')}")
        doc.add_paragraph()
        
        # Dane ucznia
        doc.add_heading('I. DANE UCZNIA', level=1)
        doc.add_paragraph(f"Imię i nazwisko: {name}")
        if birth_date:
            doc.add_paragraph(f"Data urodzenia: {birth_date}")
        if school:
            doc.add_paragraph(f"Szkoła/Przedszkole: {school}")
        if class_name:
            doc.add_paragraph(f"Klasa/Grupa: {class_name}")
        
        doc.add_paragraph()
        doc.add_heading('II. PODSTAWA PRAWNA', level=1)
        doc.add_paragraph(f"Numer orzeczenia: {cert_number}")
        if cert_date:
            doc.add_paragraph(f"Data wydania: {cert_date}")
        if valid_until:
            doc.add_paragraph(f"Ważność do: {valid_until}")
        
        # WOPFU
        doc.add_paragraph()
        doc.add_heading('III. OCENA POZIOMU FUNKCJONOWANIA (WOPFU)', level=1)
        
        doc.add_heading('Mocne strony ucznia:', level=2)
        doc.add_paragraph(self.ipet_strengths.get("1.0", "end").strip() or "Nie podano.")
        
        doc.add_heading('Trudności i ograniczenia:', level=2)
        doc.add_paragraph(self.ipet_difficulties.get("1.0", "end").strip() or "Nie podano.")
        
        doc.add_heading('Potrzeby rozwojowe i edukacyjne:', level=2)
        doc.add_paragraph(self.ipet_needs.get("1.0", "end").strip() or "Nie podano.")
        
        # Cele
        doc.add_paragraph()
        doc.add_heading('IV. CELE EDUKACYJNO-TERAPEUTYCZNE', level=1)
        
        doc.add_heading('Cele edukacyjne:', level=2)
        doc.add_paragraph(self.ipet_educational_goals.get("1.0", "end").strip() or "Nie podano.")
        
        doc.add_heading('Cele terapeutyczne:', level=2)
        doc.add_paragraph(self.ipet_therapeutic_goals.get("1.0", "end").strip() or "Nie podano.")
        
        doc.add_heading('Cele wychowawcze i społeczne:', level=2)
        doc.add_paragraph(self.ipet_social_goals.get("1.0", "end").strip() or "Nie podano.")
        
        # Dostosowania
        doc.add_paragraph()
        doc.add_heading('V. DOSTOSOWANIA I WSPARCIE', level=1)
        
        doc.add_heading('Zakres dostosowań wymagań edukacyjnych:', level=2)
        doc.add_paragraph(self.ipet_adjustments.get("1.0", "end").strip() or "Nie podano.")
        
        doc.add_heading('Formy wsparcia:', level=2)
        doc.add_paragraph(self.ipet_support_forms.get("1.0", "end").strip() or "Nie podano.")
        
        doc.add_heading('Metody i sposoby pracy:', level=2)
        doc.add_paragraph(self.ipet_teaching_methods.get("1.0", "end").strip() or "Nie podano.")
        
        # Zajęcia
        doc.add_paragraph()
        doc.add_heading('VI. ZAJĘCIA SPECJALISTYCZNE', level=1)
        doc.add_paragraph(self.ipet_therapy_sessions.get("1.0", "end").strip() or "Nie podano.")
        
        # Ewaluacja
        doc.add_paragraph()
        doc.add_heading('VII. EWALUACJA I ZESPÓŁ', level=1)
        
        doc.add_heading('Sposoby i terminy ewaluacji:', level=2)
        doc.add_paragraph(self.ipet_evaluation.get("1.0", "end").strip() or "Nie podano.")
        
        doc.add_heading('Członkowie zespołu:', level=2)
        doc.add_paragraph(self.ipet_team_members.get("1.0", "end").strip() or "Nie podano.")
        
        # Podpisy
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph("Podpis dyrektora szkoły/przedszkola")
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph("Podpisy członków zespołu")
        
        # Zapisywanie
        safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_')).replace(' ', '_')
        filename = f"IPET_{safe_name}.docx"
        filepath = os.path.join(self.save_directory, filename)
        
        try:
            doc.save(filepath)
            messagebox.showinfo("Sukces", f"Dokument IPET został zapisany:\n{filepath}")
        except PermissionError:
            messagebox.showerror("Błąd", 
                               "Brak uprawnień do zapisu w tej lokalizacji lub plik jest otwarty!\n"
                               "Spróbuj wybrać inny folder lub zamknij plik jeśli jest otwarty.")
        except Exception as e:
            messagebox.showerror("Błąd", f"Wystąpił błąd podczas zapisu:\n{str(e)}")

if __name__ == "__main__":
    app = EduApp()
    app.mainloop()
